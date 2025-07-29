"""
功能说明：
    该脚本用于自动填充 Word 模板文件，读取 Excel 数据，将数据填充到 Word 表格中，并插入相应的图片。
    
    作者：Xiongxiao Wu
    版本：1.0
    时间：20250213
"""
import pandas as pd
from docx import Document
from docx.shared import Inches
from pathlib import Path
import tkinter as tk
from tkinter import filedialog

def fill_word_template(config):
    """填充Word模板主函数"""
    try:
        df = pd.read_excel(config['excel_path'], engine='openpyxl')
    except FileNotFoundError:
        print(f"Excel文件不存在：{config['excel_path']}")
        return

    # 创建输出目录
    Path(config['output_folder']).mkdir(parents=True, exist_ok=True)

    # 图片类型配置（类型名: (行, 列, 子目录, 文件名后缀)）
    IMAGE_CONFIG = {
        '概略点位图': (4, 3, ['刺点照片', '概略点位图'], ''),
        '位置详图': (8, 3, ['刺点照片', '像控点位置详图'], ''),
        '近景': (8, 6, ['像控点点位照片', '{point_id}'], '近景'),
        '远景': (4, 6, ['像控点点位照片', '{point_id}'], '远景'),
        '中景': (6, 6, ['像控点点位照片', '{point_id}'], '中景')
    }

    for _, row in df.iterrows():
        try:
            point_id = str(row['点号'])
            doc = Document(config['word_template_path'])
            table = doc.tables[0]

            # 基础信息填充
            fill_basic_info(table, row)

            # 填充各类图片
            for img_type, (row_idx, col_idx, sub_dirs, suffix) in IMAGE_CONFIG.items():
                img_path = build_image_path(config['photo_base_path'], sub_dirs, point_id, suffix)
                insert_image_to_cell(table, img_path, row_idx, col_idx, point_id, img_type)

            # 保存文档
            output_path = Path(config['output_folder']) / f"{point_id}.docx"
            doc.save(output_path)
            print(f"成功生成：{output_path}")

        except Exception as e:
            print(f"处理点位 {point_id} 时发生严重错误：{str(e)}")


def fill_basic_info(table, row_data):
    """填充基础表格信息"""
    mappings = [
        (0, 1, '点号'),
        (1, 1, '刺点者'),
        (1, 4, '检查者'),
        (1, 8, '日期'),
        (2, 3, 'X坐标'),
        (2, 6, 'Y坐标'),
        (2, 8, 'H'),
        (9, 1, '点位说明')
    ]

    for row_idx, col_idx, field in mappings:
        cell = table.cell(row_idx, col_idx)
        cell.text = str(row_data.get(field, ''))

    # 特殊处理片号字段
    table.cell(3, 3).text = f"（片号：{row_data.get('片号', '')}）"


def build_image_path(base_path, sub_dirs, point_id, suffix):
    """构建图片路径"""
    resolved_subdirs = [d.format(point_id=point_id) for d in sub_dirs]
    filename = f"{point_id}{suffix}.jpg" if suffix else f"{point_id}.jpg"
    return Path(base_path).joinpath(*resolved_subdirs, filename)


def insert_image_to_cell(table, img_path, row_idx, col_idx, point_id, img_type):
    """向指定单元格插入图片"""
    if not img_path.exists():
        print(f"[{point_id}] {img_type}图片不存在：{img_path}")
        return

    try:
        cell = table.cell(row_idx, col_idx)
        # 清空单元格原有内容
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.clear()
        # 添加新图片
        cell.paragraphs[0].add_run().add_picture(
            str(img_path), width=Inches(3))
    except Exception as e:
        print(f"[{point_id}] 插入{img_type}图片失败：{str(e)}")


def select_data_path():
    """选择数据路径"""
    path = filedialog.askdirectory()
    if path:
        data_path_entry.delete(0, tk.END)
        data_path_entry.insert(0, path)


def select_output_folder():
    """选择输出文件夹路径"""
    path = filedialog.askdirectory()
    if path:
        output_folder_entry.delete(0, tk.END)
        output_folder_entry.insert(0, path)


def run_program():
    """运行程序"""
    data_path = data_path_entry.get()
    output_folder = output_folder_entry.get()

    # 配置信息
    CONFIG = {
        'data_path': Path(data_path),
        'output_folder': Path(output_folder),
    }

    # 派生路径配置
    CONFIG.update({
        'excel_path': CONFIG['data_path'] / "点位坐标成果.xlsx",
        'word_template_path': CONFIG['data_path'] / "点位信息模板.docx",
        'photo_base_path': CONFIG['data_path']
    })

    fill_word_template(CONFIG)


# 创建主窗口
root = tk.Tk()
root.title("Word 文件自动填充程序")

# 创建标签和输入框
tk.Label(root, text="数据路径:").grid(row=0, column=0, padx=10, pady=10)
data_path_entry = tk.Entry(root, width=50)
data_path_entry.grid(row=0, column=1, padx=10, pady=10)
tk.Button(root, text="选择数据路径", command=select_data_path).grid(row=0, column=2, padx=10, pady=10)

tk.Label(root, text="输出文件夹路径:").grid(row=1, column=0, padx=10, pady=10)
output_folder_entry = tk.Entry(root, width=50)
output_folder_entry.grid(row=1, column=1, padx=10, pady=10)
tk.Button(root, text="选择输出文件夹", command=select_output_folder).grid(row=1, column=2, padx=10, pady=10)

# 创建运行按钮
tk.Button(root, text="运行", command=run_program).grid(row=2, column=1, padx=10, pady=20)

# 运行主循环
root.mainloop()