import pandas as pd
from docx import Document
from docx.shared import Inches
from pathlib import Path


def fill_word_template(config):
    """填充Word模板主函数"""
    try:
        df = pd.read_excel(config['excel_path'], engine='openpyxl')
    except FileNotFoundError:
        print(f"Excel文件不存在：{config['excel_path']}")
        return

    # 创建输出目录
    Path(config['output_folder']).mkdir(parents=True, exist_ok=True)

    # 图片类型配置（类型名: (行, 列, 子目录, 文件名后缀)）
    IMAGE_CONFIG = {
        '概略点位图': (4, 3, ['刺点照片', '概略点位图'], ''),
        '位置详图': (8, 3, ['刺点照片', '像控点位置详图'], ''),
        '近景': (8, 6, ['像控点点位照片', '{point_id}'], '近景'),
        '远景': (4, 6, ['像控点点位照片', '{point_id}'], '远景'),
        '中景': (6, 6, ['像控点点位照片', '{point_id}'], '中景')
    }

    for _, row in df.iterrows():
        try:
            point_id = str(row['点号'])
            doc = Document(config['word_template_path'])
            table = doc.tables[0]

            # 基础信息填充
            fill_basic_info(table, row)

            # 填充各类图片
            for img_type, (row_idx, col_idx, sub_dirs, suffix) in IMAGE_CONFIG.items():
                img_path = build_image_path(config['photo_base_path'], sub_dirs, point_id, suffix)
                insert_image_to_cell(table, img_path, row_idx, col_idx, point_id, img_type)

            # 保存文档
            output_path = Path(config['output_folder']) / f"{point_id}.docx"
            doc.save(output_path)
            print(f"成功生成：{output_path}")

        except Exception as e:
            print(f"处理点位 {point_id} 时发生严重错误：{str(e)}")


def fill_basic_info(table, row_data):
    """填充基础表格信息"""
    mappings = [
        (0, 1, '点号'),
        (1, 1, '刺点者'),
        (1, 4, '检查者'),
        (1, 8, '日期'),
        (2, 3, 'X坐标'),
        (2, 6, 'Y坐标'),
        (2, 8, 'H'),
        (9, 1, '点位说明')
    ]

    for row_idx, col_idx, field in mappings:
        cell = table.cell(row_idx, col_idx)
        cell.text = str(row_data.get(field, ''))

    # 特殊处理片号字段
    table.cell(3, 3).text = f"（片号：{row_data.get('片号', '')}）"


def build_image_path(base_path, sub_dirs, point_id, suffix):
    """构建图片路径"""
    resolved_subdirs = [d.format(point_id=point_id) for d in sub_dirs]
    filename = f"{point_id}{suffix}.jpg" if suffix else f"{point_id}.jpg"
    return Path(base_path).joinpath(*resolved_subdirs, filename)


def insert_image_to_cell(table, img_path, row_idx, col_idx, point_id, img_type):
    """向指定单元格插入图片"""
    if not img_path.exists():
        print(f"[{point_id}] {img_type}图片不存在：{img_path}")
        return

    try:
        cell = table.cell(row_idx, col_idx)
        # 清空单元格原有内容
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.clear()
        # 添加新图片
        cell.paragraphs[0].add_run().add_picture(
            str(img_path), width=Inches(3))
    except Exception as e:
        print(f"[{point_id}] 插入{img_type}图片失败：{str(e)}")


if __name__ == "__main__":
    # 功能说明：
    # 该脚本用于自动填充 Word 模板文件，读取 Excel 数据，将数据填充到 Word 表格中，并插入相应的图片。
    # 使用说明：
    # 1. 修改 CONFIG 字典中的 data_path 和 output_folder 路径信息。
    # 2. 确保 Python 环境中已经安装了 pandas 和 python-docx 库。
    # 3. 运行该 Python 脚本。
    # 作者：Wu Xiongxiao
    # 版本：1.0
    #时间：20250213


    # 配置信息
    CONFIG = {
        'data_path': Path("/Users/WXX/Documents/同步空间/1.学习脚本/python/像控点位信息表自动填充/样例数据"),
        'output_folder': Path("/Users/WXX/Desktop/点位信息表"),
    }

    # 派生路径配置
    CONFIG.update({
        'excel_path': CONFIG['data_path'] / "点位坐标成果.xlsx",
        'word_template_path': CONFIG['data_path'] / "点位信息模板.docx",
        'photo_base_path': CONFIG['data_path']
    })

    fill_word_template(CONFIG)